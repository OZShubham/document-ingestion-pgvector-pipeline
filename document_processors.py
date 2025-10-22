from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from typing import Dict, List, Any, Optional
import io
import logging
 
logger = logging.getLogger(__name__)
 
@dataclass
class ProcessedDocument:
    """Standardized document processing result"""
    text: str
    metadata: Dict[str, Any]
    sections: List[Dict[str, Any]] = field(default_factory=list)
    tables: List[Dict[str, Any]] = field(default_factory=list)
    images: List[Dict[str, Any]] = field(default_factory=list)
    page_count: int = 0
    processing_method: str = ""
    error: Optional[str] = None
 
class BaseDocumentProcessor(ABC):
    """Base class for document processors"""
   
    @abstractmethod
    async def process(self, file_bytes: bytes, filename: str, **kwargs) -> ProcessedDocument:
        pass
   
    @abstractmethod
    def supports(self, mime_type: str) -> bool:
        pass
 
class GeminiProcessor(BaseDocumentProcessor):
    """Process documents using Gemini"""
   
    def __init__(self):
        from google import genai
        from config import Config
        self.client = genai.Client(
            vertexai=True,
            project=Config.PROJECT_ID,
            location=Config.REGION
        )
        self.config = Config
   
    def supports(self, mime_type: str) -> bool:
        return mime_type in self.config.GEMINI_SUPPORTED_TYPES
   
    async def process(self, file_bytes: bytes, filename: str, mime_type: str = None, **kwargs) -> ProcessedDocument:
        from google.genai.types import Part, GenerateContentConfig
        import json
       
        extraction_prompt = """Analyze this document and extract:
        1. Full text content preserving structure
        2. All tables in markdown format
        3. Descriptions of images, charts, diagrams
        4. Document metadata (title, author, date, pages if present)
        5. Logical sections with headings
       
        Return as JSON:
        {
            "text": "full text",
            "metadata": {"title": "", "author": "", "date": "", "pages": 0},
            "sections": [{"heading": "", "content": "", "page": 0}],
            "tables": [{"content": "", "description": "", "page": 0}],
            "images": [{"description": "", "page": 0}]
        }"""
       
        try:
            response = await self.client.aio.models.generate_content(
                model=self.config.GEMINI_MODEL,
                contents=[
                    extraction_prompt,
                    Part.from_bytes(data=file_bytes, mime_type=mime_type)
                ],
                config=GenerateContentConfig(
                    temperature=0,
                    response_mime_type="application/json"
                )
            )
           
            result = json.loads(response.text)
           
            return ProcessedDocument(
                text=result.get('text', ''),
                metadata=result.get('metadata', {}),
                sections=result.get('sections', []),
                tables=result.get('tables', []),
                images=result.get('images', []),
                page_count=result.get('metadata', {}).get('pages', 0),
                processing_method='gemini'
            )
           
        except Exception as e:
            logger.error(f"Gemini processing failed: {e}")
            return ProcessedDocument(
                text="",
                metadata={},
                processing_method='gemini',
                error=str(e)
            )
 
class PyMuPDFProcessor(BaseDocumentProcessor):
    """Process PDFs using PyMuPDF"""
   
    def supports(self, mime_type: str) -> bool:
        return mime_type == 'application/pdf'
   
    async def process(self, file_bytes: bytes, filename: str, **kwargs) -> ProcessedDocument:
        import fitz
       
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
           
            full_text = []
            tables = []
            images = []
           
            for page_num, page in enumerate(doc, start=1):
                text = page.get_text()
                full_text.append(text)
               
                # Extract images
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    images.append({
                        'page': page_num,
                        'index': img_index,
                        'description': f'Image {img_index + 1} on page {page_num}'
                    })
               
                # Extract tables (if supported)
                try:
                    if hasattr(page, 'find_tables'):
                        tables_on_page = page.find_tables()
                        if tables_on_page:
                            for table_index, table in enumerate(tables_on_page):
                                try:
                                    tables.append({
                                        'page': page_num,
                                        'content': table.extract(),
                                        'description': f'Table {table_index + 1} on page {page_num}'
                                    })
                                except:
                                    pass
                except Exception:
                    pass
           
            return ProcessedDocument(
                text="\n\n".join(full_text),
                metadata={
                    'title': doc.metadata.get('title', ''),
                    'author': doc.metadata.get('author', ''),
                    'pages': len(doc)
                },
                tables=tables,
                images=images,
                page_count=len(doc),
                processing_method='pymupdf'
            )
           
        except Exception as e:
            logger.error(f"PyMuPDF processing failed: {e}")
            return ProcessedDocument(
                text="",
                metadata={},
                processing_method='pymupdf',
                error=str(e)
            )
 
class PyPDFProcessor(BaseDocumentProcessor):
    """Process PDFs using PyPDF"""
   
    def supports(self, mime_type: str) -> bool:
        return mime_type == 'application/pdf'
   
    async def process(self, file_bytes: bytes, filename: str, **kwargs) -> ProcessedDocument:
        from pypdf import PdfReader
       
        try:
            pdf = PdfReader(io.BytesIO(file_bytes))
           
            full_text = []
            for page in pdf.pages:
                full_text.append(page.extract_text())
           
            metadata = pdf.metadata or {}
           
            return ProcessedDocument(
                text="\n\n".join(full_text),
                metadata={
                    'title': metadata.get('/Title', ''),
                    'author': metadata.get('/Author', ''),
                    'pages': len(pdf.pages)
                },
                page_count=len(pdf.pages),
                processing_method='pypdf'
            )
           
        except Exception as e:
            logger.error(f"PyPDF processing failed: {e}")
            return ProcessedDocument(
                text="",
                metadata={},
                processing_method='pypdf',
                error=str(e)
            )
 
class DocxProcessor(BaseDocumentProcessor):
    """Process Word documents"""
   
    def supports(self, mime_type: str) -> bool:
        return mime_type in {
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword'
        }
   
    async def process(self, file_bytes: bytes, filename: str, **kwargs) -> ProcessedDocument:
        from docx import Document
       
        try:
            doc = Document(io.BytesIO(file_bytes))
           
            full_text = []
            tables = []
           
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
           
            for table_index, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
               
                tables.append({
                    'content': table_data,
                    'description': f'Table {table_index + 1}'
                })
           
            core_props = doc.core_properties
           
            return ProcessedDocument(
                text="\n\n".join(full_text),
                metadata={
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'created': str(core_props.created) if core_props.created else ''
                },
                tables=tables,
                processing_method='docx'
            )
           
        except Exception as e:
            logger.error(f"DOCX processing failed: {e}")
            return ProcessedDocument(
                text="",
                metadata={},
                processing_method='docx',
                error=str(e)
            )
 
class ExcelProcessor(BaseDocumentProcessor):
    """Process Excel files"""
   
    def supports(self, mime_type: str) -> bool:
        return mime_type in {
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'application/vnd.ms-excel'
        }
   
    async def process(self, file_bytes: bytes, filename: str, **kwargs) -> ProcessedDocument:
        from openpyxl import load_workbook
       
        try:
            wb = load_workbook(io.BytesIO(file_bytes), read_only=True)
           
            full_text = []
            tables = []
           
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
               
                sheet_data = []
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        sheet_data.append([str(cell) if cell is not None else '' for cell in row])
               
                if sheet_data:
                    text_rows = [' | '.join(row) for row in sheet_data]
                    full_text.append(f"Sheet: {sheet_name}\n" + '\n'.join(text_rows))
                   
                    tables.append({
                        'content': sheet_data,
                        'description': f'Sheet: {sheet_name}'
                    })
           
            return ProcessedDocument(
                text="\n\n".join(full_text),
                metadata={'sheets': wb.sheetnames},
                tables=tables,
                processing_method='openpyxl'
            )
           
        except Exception as e:
            logger.error(f"Excel processing failed: {e}")
            return ProcessedDocument(
                text="",
                metadata={},
                processing_method='openpyxl',
                error=str(e)
            )
 
class TextProcessor(BaseDocumentProcessor):
    """Process plain text files"""
   
    def supports(self, mime_type: str) -> bool:
        return mime_type in {'text/plain', 'text/markdown', 'text/csv'}
   
    async def process(self, file_bytes: bytes, filename: str, **kwargs) -> ProcessedDocument:
        try:
            text = file_bytes.decode('utf-8')
            return ProcessedDocument(
                text=text,
                metadata={'encoding': 'utf-8'},
                processing_method='text'
            )
        except UnicodeDecodeError:
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    text = file_bytes.decode(encoding)
                    return ProcessedDocument(
                        text=text,
                        metadata={'encoding': encoding},
                        processing_method='text'
                    )
                except:
                    continue
           
            return ProcessedDocument(
                text="",
                metadata={},
                processing_method='text',
                error="Unable to decode text file"
            )
 
class DocumentProcessorFactory:
    """Factory to get appropriate document processor"""
   
    def __init__(self):
        self.processors = {
            'gemini': GeminiProcessor(),
            'pymupdf': PyMuPDFProcessor(),
            'pypdf': PyPDFProcessor(),
            'docx': DocxProcessor(),
            'openpyxl': ExcelProcessor(),
            'text': TextProcessor(),
        }
   
    async def process_document(
        self,
        file_bytes: bytes,
        filename: str,
        mime_type: str,
        preferred_method: str = None
    ) -> ProcessedDocument:
        """Process document with fallback strategies"""
        from config import Config
       
        available_methods = Config.PROCESSING_METHODS.get(mime_type, [])
       
        if not available_methods:
            return ProcessedDocument(
                text="",
                metadata={},
                error=f"Unsupported file type: {mime_type}"
            )
       
        if preferred_method and preferred_method in available_methods:
            methods_to_try = [preferred_method] + [m for m in available_methods if m != preferred_method]
        else:
            methods_to_try = available_methods
       
        last_error = None
       
        for method in methods_to_try:
            processor = self.processors.get(method)
            if processor and processor.supports(mime_type):
                result = await processor.process(
                    file_bytes=file_bytes,
                    filename=filename,
                    mime_type=mime_type
                )
               
                if result.error is None and result.text:
                    return result
               
                last_error = result.error
       
        return ProcessedDocument(
            text="",
            metadata={},
            error=f"All processing methods failed. Last error: {last_error}"
        )